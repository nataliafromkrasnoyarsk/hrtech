#!/usr/bin/env python3
"""Convert the article markdown to a formatted .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

INPUT = "/home/user/hrtech/content-projects/output/drafts/osipova-management-2026/2026-03-13-management-practices-2026-draft-v1.md"
OUTPUT = "/home/user/hrtech/content-projects/output/drafts/osipova-management-2026/2026-03-13-management-practices-2026-draft-v1.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.15

# Heading styles
for level, size, color in [(1, 18, '1a1a1a'), (2, 15, '1a1a1a'), (3, 13, '333333')]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(*bytes.fromhex(color))
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    hs.paragraph_format.space_after = Pt(8)

with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()


def add_rich_paragraph(doc, text, style_name='Normal', is_italic_block=False):
    """Add a paragraph with bold/italic markdown formatting."""
    p = doc.add_paragraph(style=style_name)
    if is_italic_block:
        # Whole paragraph is italic (wrapped in * *)
        text = text.strip('*').strip()
        run = p.add_run(text)
        run.italic = True
        return p

    # Process **bold** and *italic* inline
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p


i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Headings
    if line.startswith('### '):
        doc.add_heading(line[4:], level=3)
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        add_rich_paragraph(doc, m.group(2), 'List Number' if 'List Number' in [s.name for s in doc.styles] else 'Normal')
        i += 1
        continue

    # Bullet list
    if line.startswith('- '):
        add_rich_paragraph(doc, line[2:], 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else 'Normal')
        i += 1
        continue

    # Italic block (author line, bio)
    if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
        add_rich_paragraph(doc, line, is_italic_block=True)
        i += 1
        continue

    # Regular paragraph
    add_rich_paragraph(doc, line)
    i += 1

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
